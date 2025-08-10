import os
import re
from abc import ABC, abstractmethod
from bs4 import BeautifulSoup
from typing import Dict, List, Any, Union, Optional
import requests
from youtube_transcript_api import YouTubeTranscriptApi
import fitz  # PyMuPDF
import docx
import urllib.parse
import json

from app.core.config import settings


class DocumentProcessor(ABC):
    """Base class for document processors."""
    
    @abstractmethod
    def process(self, source: str, content: Optional[bytes] = None) -> List[Dict[str, Any]]:
        """
        Process a document and return chunks with metadata.
        
        Args:
            source: Source identifier (URL or file path)
            content: Optional binary content of the document
            
        Returns:
            List of dictionaries containing text chunks and metadata
        """
        pass
    
    def chunk_text(self, text: str, metadata: Dict[str, Any]) -> List[Dict[str, Any]]:
        """
        Split text into chunks with overlap using a more context-aware approach.
        
        Args:
            text: Text to chunk
            metadata: Metadata to include with each chunk
            
        Returns:
            List of dictionaries with text and metadata
        """
        # Clean the text but preserve more meaningful whitespace patterns
        text = re.sub(r'\s{3,}', '\n\n', text)  # Convert excessive whitespace to paragraph breaks
        text = re.sub(r'\s{2}', '\n', text)     # Convert double spaces to line breaks
        text = re.sub(r'\t', ' ', text)         # Replace tabs with spaces
        text = text.strip()
        
        if not text:
            print(f"Warning: Empty text content for source {metadata.get('source', 'unknown')}")
            return []
        
        # For PDFs, try to maintain page boundary information if available
        page_markers = []
        if metadata.get('source_type') == 'pdf':
            # Capture page markers like [Page 1], [Page 2] etc.
            page_pattern = re.compile(r'\[Page \d+\]')
            page_markers = [(m.start(), m.end()) for m in page_pattern.finditer(text)]
        
        # Short text handling - if the text is very small, don't chunk it
        if len(text) < 200:  # Increased small size threshold
            print(f"Text content is small ({len(text)} chars), not chunking")
            chunk_metadata = metadata.copy()
            chunk_metadata.update({
                "chunk_index": 0,
                "char_start": 0,
                "char_end": len(text)
            })
            return [{
                "text": text,
                "metadata": chunk_metadata
            }]
        
        # Use a more accurate token estimation (especially for non-English text)
        chars_per_token = 3.5  # Better approximation for mixed content
        # Adjust chunk size based on document type for better context retention
        base_chunk_size = settings.chunk_size
        if metadata.get('source_type') == 'pdf':
            # PDF documents may need larger chunks to maintain context
            base_chunk_size = int(base_chunk_size * 1.2)
        
        # Calculate chunk sizes and overlap
        chunk_size_chars = base_chunk_size * chars_per_token
        overlap_chars = int(chunk_size_chars * settings.chunk_overlap)
        
        # Use recursive chunking for better semantic cohesion
        chunks = []
        headings = []
        
        # Try to identify headings or sections using common patterns
        heading_patterns = [
            r'^#{1,6}\s+.+$',                 # Markdown headings
            r'^[A-Z][A-Za-z\s]{2,50}$',       # Capitalized section titles
            r'^\d+(\.\d+)*\s+[A-Z]',          # Numbered sections
            r'^[IVXLCDMivxlcdm]+\.\s+.+$'     # Roman numeral sections
        ]
        
        # Find potential section breaks
        for pattern in heading_patterns:
            for match in re.finditer(pattern, text, re.MULTILINE):
                headings.append(match.start())
        
        # Sort headings for proper chunking
        headings.sort()
        
        # Function for recursive chunking that respects section boundaries
        def create_chunks(start_pos, end_pos, depth=0):
            # Ensure start_pos and end_pos are integers
            start_pos = int(start_pos)
            end_pos = int(end_pos)
            
            section_text = text[start_pos:end_pos].strip()
            section_length = len(section_text)
            
            # Base case: if section is small enough for one chunk
            if section_length <= chunk_size_chars:
                if section_text:
                    chunk_metadata = metadata.copy()
                    chunk_metadata.update({
                        "chunk_index": len(chunks),
                        "char_start": start_pos,
                        "char_end": end_pos,
                        "depth": depth
                    })
                    
                    # Add page number info if available
                    for p_start, p_end in page_markers:
                        if start_pos <= p_start < end_pos:
                            page_info = text[p_start:p_end]
                            chunk_metadata["page_info"] = page_info
                            break
                    
                    chunks.append({
                        "text": section_text,
                        "metadata": chunk_metadata
                    })
                return
            
            # Find the best breaking point in this section
            mid_point = start_pos + section_length // 2
            best_break = mid_point
            
            # Try to find paragraph breaks near the middle
            # Handle potential floating-point values in chunk_size_chars
            half_chunk = int(chunk_size_chars//2)
            quarter_chunk = int(chunk_size_chars//4)
            
            # Make sure all arguments are integers
            para_break_pos_nn = text.rfind('\n\n', start_pos, mid_point + half_chunk)
            para_break_pos_rnrn = text.rfind('\r\n\r\n', start_pos, mid_point + half_chunk)
            
            # Use -1 as the default value when not found instead of relying on max()
            paragraph_break = max(para_break_pos_nn, para_break_pos_rnrn)
            
            if paragraph_break > start_pos:
                best_break = paragraph_break
            else:
                # Try sentence breaks
                sent_break_period = text.rfind('. ', start_pos, mid_point + quarter_chunk)
                sent_break_excl = text.rfind('! ', start_pos, mid_point + quarter_chunk)
                sent_break_quest = text.rfind('? ', start_pos, mid_point + quarter_chunk)
                
                # Find the max of valid break points (-1 means not found)
                valid_breaks = [b for b in [sent_break_period, sent_break_excl, sent_break_quest] if b >= 0]
                sentence_break = max(valid_breaks) if valid_breaks else -1
                if sentence_break > start_pos:
                    best_break = sentence_break + 1  # Include the punctuation
            
            # Recursively chunk the two halves
            create_chunks(start_pos, best_break, depth + 1)
            
            # Add overlap for context continuity
            overlap_start = max(start_pos, best_break - overlap_chars)
            create_chunks(overlap_start, end_pos, depth + 1)
        
        # Begin the recursive chunking process
        create_chunks(0, len(text))
        
        # Sort chunks by position and deduplicate
        chunks.sort(key=lambda x: x['metadata']['char_start'])
        
        # Remove duplicate or highly overlapping chunks
        deduped_chunks = []
        for i, chunk in enumerate(chunks):
            if i == 0 or chunk['text'] != chunks[i-1]['text']:
                deduped_chunks.append(chunk)
        
        # Renumber chunk indices
        for i, chunk in enumerate(deduped_chunks):
            chunk['metadata']['chunk_index'] = i
        
        print(f"Created {len(deduped_chunks)} semantic chunks from {len(text)} characters")
        return deduped_chunks


class PDFProcessor(DocumentProcessor):
    """Enhanced processor for PDF documents with better structure preservation."""
    
    def process(self, source: str, content: Optional[bytes] = None) -> List[Dict[str, Any]]:
        try:
            # Create base metadata
            metadata = {
                "source": source,
                "source_type": "pdf",
            }
            
            print(f"Starting PDF processing: {source}")
            start_time = __import__('time').time()
            
            # Open PDF from content or file
            if content:
                doc = fitz.open(stream=content, filetype="pdf")
            else:
                doc = fitz.open(source)
            
            # Extract metadata
            pdf_metadata = doc.metadata
            if pdf_metadata:
                metadata.update({
                    "title": pdf_metadata.get("title", "") or os.path.basename(source),
                    "author": pdf_metadata.get("author", ""),
                    "creation_date": pdf_metadata.get("creationDate", ""),
                    "subject": pdf_metadata.get("subject", ""),
                    "keywords": pdf_metadata.get("keywords", ""),
                })
                
            # Add page count to metadata
            page_count = len(doc)
            metadata["page_count"] = page_count
            
            print(f"PDF has {page_count} pages")
            
            # Extract table of contents if available
            toc = doc.get_toc()
            if toc:
                metadata["toc"] = toc
                print(f"Found table of contents with {len(toc)} entries")
            
            # Process in batches for better performance
            all_text = []
            all_headings = []
            batch_size = min(20, page_count)  # Process up to 20 pages at once
            
            # Extract bookmarks for section detection
            bookmarks = {item[2]: item[1] for item in toc} if toc else {}
            
            for batch_start in range(0, page_count, batch_size):
                batch_end = min(batch_start + batch_size, page_count)
                batch_text = []
                
                # Process a batch of pages
                for page_num in range(batch_start, batch_end):
                    page = doc[page_num]
                    
                    # Check if page has a bookmark/section heading
                    section_heading = None
                    if page_num + 1 in bookmarks:
                        section_heading = bookmarks[page_num + 1]
                    
                    # Try to extract structured content - first HTML to preserve more layout info
                    try:
                        # First try HTML extraction for better structure
                        html_text = page.get_text("html")
                        if html_text:
                            # Parse HTML to extract structured content
                            from bs4 import BeautifulSoup
                            soup = BeautifulSoup(html_text, 'html.parser')
                            
                            # Extract text with better formatting
                            page_content = []
                            
                            # Extract fonts and styles to identify headers
                            fonts = {}
                            for span in soup.find_all('span'):
                                if 'style' in span.attrs:
                                    style = span.get('style', '')
                                    size_match = re.search(r'font-size:(\d+)px', style)
                                    weight_match = re.search(r'font-weight:(\d+)', style)
                                    
                                    if size_match:
                                        size = int(size_match.group(1))
                                        if size > 14:  # Likely heading
                                            # Format as heading
                                            heading_text = span.get_text().strip()
                                            if heading_text and len(heading_text) < 100:  # Reasonable heading length
                                                page_content.append(f"\n## {heading_text}\n")
                                                continue
                            
                            # Get cleaned page text
                            clean_text = soup.get_text()
                            if clean_text:
                                page_content.append(clean_text)
                            
                            # Add section heading if found
                            if section_heading:
                                page_text = f"[Section: {section_heading}]\n" + "\n".join(page_content)
                            else:
                                page_text = "\n".join(page_content)
                    except:
                        # Fallback to simple text extraction
                        page_text = page.get_text("text")
                    
                    # Extract blocks for better layout understanding
                    blocks = page.get_text("blocks")
                    if blocks:
                        # Try to identify headings based on font size or position
                        for b in blocks:
                            if b[5] > 1.5 * 12:  # Font size much larger than normal text
                                heading_text = b[4].strip()
                                if heading_text and len(heading_text) < 100:  # Reasonable heading length
                                    all_headings.append(heading_text)
                    
                    # Add page number and any sections to the text
                    batch_text.append(f"[Page {page_num + 1}]\n{page_text}")
                
                all_text.extend(batch_text)
                print(f"Processed pages {batch_start+1} to {batch_end} of {page_count}")
            
            # Close the document
            doc.close()
            
            # Add identified headings to metadata for better context
            if all_headings:
                metadata["headings"] = all_headings[:20]  # Store up to 20 headings
            
            # Join all text
            complete_text = "\n\n".join(all_text)
            
            # Enhanced text cleanup that preserves important structure
            complete_text = re.sub(r'[ \t]+', ' ', complete_text)
            complete_text = re.sub(r'\n{4,}', '\n\n', complete_text)
            complete_text = re.sub(r'(\[Page \d+\])\s+(\[Page \d+\])', r'\1\n\n\2', complete_text)
            
            end_time = __import__('time').time()
            processing_time = end_time - start_time
            print(f"PDF processing completed in {processing_time:.2f} seconds")
            print(f"Extracted {len(complete_text)} characters")
            
            # Chunk the text with enhanced metadata
            return self.chunk_text(complete_text, metadata)
        
        except Exception as e:
            print(f"Error processing PDF {source}: {e}")
            import traceback
            traceback.print_exc()
            return []


class DocxProcessor(DocumentProcessor):
    """Processor for DOCX documents."""
    
    def process(self, source: str, content: Optional[bytes] = None) -> List[Dict[str, Any]]:
        try:
            # Create base metadata
            metadata = {
                "source": source,
                "source_type": "docx",
            }
            
            # Open DOCX from content or file
            if content:
                import io
                doc = docx.Document(io.BytesIO(content))
            else:
                doc = docx.Document(source)
            
            # Extract text from paragraphs
            all_text = "\n".join([para.text for para in doc.paragraphs if para.text])
            
            # Extract titles and headers if available
            if doc.core_properties.title:
                metadata["title"] = doc.core_properties.title
            
            # Chunk the text
            return self.chunk_text(all_text, metadata)
        
        except Exception as e:
            print(f"Error processing DOCX {source}: {e}")
            return []


class WebPageProcessor(DocumentProcessor):
    """Processor for web pages."""
    
    def process(self, source: str, content: Optional[bytes] = None) -> List[Dict[str, Any]]:
        try:
            # Create base metadata
            metadata = {
                "source": source,
                "source_type": "web",
            }
            
            # Check if source is actually a URL
            if not source.startswith(('http://', 'https://')):
                raise ValueError(f"Invalid URL format: {source}. Must start with http:// or https://")
                
            print(f"Processing web page: {source}")
            
            # Get HTML content with timeout and headers for better performance
            if content:
                html_content = content.decode('utf-8', errors='replace')
                print(f"Using provided content for {source}")
            else:
                # Use a session for connection pooling
                session = requests.Session()
                headers = {
                    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36",
                    "Accept": "text/html,application/xhtml+xml,application/xml",
                    "Accept-Language": "en-US,en;q=0.9",
                }
                
                print(f"Fetching content from URL: {source}")
                # Use a shorter timeout for faster response
                response = session.get(source, headers=headers, timeout=10)
                response.raise_for_status()
                html_content = response.text
                print(f"Successfully fetched {len(html_content)} bytes from {source}")
            
            # Try to use lxml parser for faster HTML parsing, fallback to html.parser
            try:
                soup = BeautifulSoup(html_content, 'lxml')
            except:
                # Fallback to built-in parser if lxml isn't available
                soup = BeautifulSoup(html_content, 'html.parser')
                print("Using fallback HTML parser")
            
            # Extract title and metadata
            if soup.title:
                title = soup.title.text.strip()
                metadata["title"] = title
                print(f"Found page title: {title}")
            
            # Extract meta description if available
            meta_desc = soup.find('meta', attrs={'name': 'description'})
            if meta_desc and 'content' in meta_desc.attrs:
                metadata["description"] = meta_desc['content']
                print(f"Found meta description: {metadata['description'][:50]}...")
            
            # Extract canonical URL if available
            canonical = soup.find('link', attrs={'rel': 'canonical'})
            if canonical and 'href' in canonical.attrs:
                metadata["canonical_url"] = canonical['href']
            
            print("Cleaning HTML content...")
            
            # Remove unwanted elements more efficiently (one pass)
            for element in soup(["script", "style", "nav", "footer", "header", "aside", "iframe", "noscript", "svg"]):
                element.decompose()
            
            # Try to identify content more intelligently
            content_elements = []
            
            # Look for common content containers
            main_content = (
                soup.find(id=re.compile(r'(content|main|article|post)', re.I)) or
                soup.find("main") or 
                soup.find("article") or 
                soup.find(attrs={"class": re.compile(r'(content|main|article|post)', re.I)}) or
                soup.find("body")
            )
            
            if main_content:
                # Extract text from semantic elements with priority
                for tag in ["h1", "h2", "h3", "p", "li", "h4", "h5", "h6", "blockquote"]:
                    elements = main_content.find_all(tag)
                    for el in elements:
                        if el.text.strip() and len(el.text.strip()) > 20:  # Ignore short elements (likely UI text)
                            content_elements.append(el.text.strip())
            else:
                # Fallback: extract all paragraphs
                for p in soup.find_all("p"):
                    if p.text.strip() and len(p.text.strip()) > 20:
                        content_elements.append(p.text.strip())
            
            # If we still don't have content, use body text
            if not content_elements:
                all_text = soup.get_text()
            else:
                all_text = "\n\n".join(content_elements)
            
            # Clean text more aggressively
            all_text = re.sub(r'\s+', ' ', all_text).strip()
            all_text = re.sub(r'\s*\.\s*', '. ', all_text)  # Fix spacing around periods
            
            print(f"Extracted {len(all_text)} characters from {source}")
            
            # Chunk the text
            return self.chunk_text(all_text, metadata)
        
        except Exception as e:
            print(f"Error processing web page {source}: {e}")
            return []


class YouTubeProcessor(DocumentProcessor):
    """Processor for YouTube videos."""
    
    @staticmethod
    def extract_video_id(url: str) -> Optional[str]:
        """Extract video ID from a YouTube URL."""
        # Match formats: youtu.be/xxx, youtube.com/watch?v=xxx, youtube.com/v/xxx
        patterns = [
            r'(?:youtu\.be/|youtube\.com/(?:embed/|v/|watch\?v=|watch\?.+&v=))([^?&/]+)',
            r'youtube\.com/shorts/([^?&/]+)'
        ]
        
        for pattern in patterns:
            match = re.search(pattern, url)
            if match:
                return match.group(1)
        return None
    
    def process(self, source: str, content: Optional[bytes] = None) -> List[Dict[str, Any]]:
        try:
            # Extract video ID from URL
            video_id = self.extract_video_id(source)
            if not video_id:
                raise ValueError(f"Could not extract video ID from URL: {source}")
            
            # Create base metadata
            metadata = {
                "source": source,
                "source_type": "youtube",
                "video_id": video_id,
            }
            
            # If content is provided, assume it's a pre-fetched transcript
            if content:
                transcript_text = content.decode('utf-8', errors='replace')
            else:
                # Get transcript
                transcript_list = YouTubeTranscriptApi.get_transcript(video_id)
                
                # Format transcript
                transcript_parts = []
                for item in transcript_list:
                    text = item.get('text', '').strip()
                    start = item.get('start', 0)
                    duration = item.get('duration', 0)
                    
                    if text:
                        timestamp = self._format_timestamp(start)
                        transcript_parts.append(f"[{timestamp}] {text}")
                
                transcript_text = "\n".join(transcript_parts)
            
            # Get video metadata
            # We're using a simplified approach here to avoid API keys
            try:
                # Try to get basic info from oEmbed
                oembed_url = f"https://www.youtube.com/oembed?url=https://www.youtube.com/watch?v={video_id}&format=json"
                response = requests.get(oembed_url)
                if response.status_code == 200:
                    video_data = response.json()
                    metadata["title"] = video_data.get("title", "")
                    metadata["author"] = video_data.get("author_name", "")
            except Exception:
                # Fallback - don't fail processing if we can't get video metadata
                pass
            
            # Chunk the transcript
            return self.chunk_text(transcript_text, metadata)
        
        except Exception as e:
            print(f"Error processing YouTube video {source}: {e}")
            return []
    
    @staticmethod
    def _format_timestamp(seconds: float) -> str:
        """Format seconds into a MM:SS timestamp."""
        minutes, seconds = divmod(int(seconds), 60)
        return f"{minutes:02d}:{seconds:02d}"


def get_processor(source: str) -> DocumentProcessor:
    """
    Get the appropriate processor for a given source.
    
    Args:
        source: Source URL or file path
        
    Returns:
        Appropriate DocumentProcessor instance
    """
    # Check if source is a local file path or temp file
    is_local_file = (
        source.startswith('/') or
        source.startswith('C:\\') or
        source.startswith('c:\\') or
        source.startswith('D:\\') or
        source.startswith('d:\\') or
        '\\AppData\\Local\\Temp\\' in source or
        '/tmp/' in source
    )
    
    # Process based on file extension or URL type
    if source.endswith('.pdf'):
        return PDFProcessor()
    elif source.endswith('.docx'):
        return DocxProcessor()
    elif source.endswith('.txt') and is_local_file:
        # Create a more robust text processor for text files
        class TextFileProcessor(DocumentProcessor):
            def process(self, source: str, content: Optional[bytes] = None) -> List[Dict[str, Any]]:
                try:
                    # Extract filename from path for better metadata
                    import os
                    filename = os.path.basename(source)
                    
                    metadata = {
                        "source": source, 
                        "source_type": "text",
                        "filename": filename
                    }
                    
                    print(f"Processing text file: {source}")
                    
                    if content:
                        print(f"Using provided content for {source}")
                        # Try multiple encodings if the first one fails
                        encodings = ['utf-8', 'latin-1', 'cp1252']
                        text = None
                        
                        for encoding in encodings:
                            try:
                                text = content.decode(encoding, errors='replace')
                                print(f"Successfully decoded with {encoding} encoding")
                                break
                            except UnicodeDecodeError:
                                continue
                        
                        if text is None:
                            # Fallback to utf-8 with replace if all else fails
                            text = content.decode('utf-8', errors='replace')
                    else:
                        print(f"Reading text from file: {source}")
                        # Try multiple encodings
                        for encoding in ['utf-8', 'latin-1', 'cp1252']:
                            try:
                                with open(source, 'r', encoding=encoding) as f:
                                    text = f.read()
                                print(f"Successfully read file with {encoding} encoding")
                                break
                            except (UnicodeDecodeError, FileNotFoundError) as e:
                                if isinstance(e, FileNotFoundError):
                                    raise  # Re-raise if file not found
                                continue
                        else:
                            # Fallback if all encodings fail
                            with open(source, 'r', encoding='utf-8', errors='replace') as f:
                                text = f.read()
                    
                    # Enhanced metadata
                    metadata["char_count"] = len(text)
                    metadata["line_count"] = text.count('\n') + 1
                    
                    print(f"Text file processing complete: {len(text)} characters, {metadata['line_count']} lines")
                    
                    return self.chunk_text(text, metadata)
                except Exception as e:
                    print(f"Error processing text file {source}: {e}")
                    return []
                    
        return TextFileProcessor()
    elif 'youtube.com' in source or 'youtu.be' in source:
        return YouTubeProcessor()
    elif source.startswith('http://') or source.startswith('https://'):
        # Only use web processor for actual URLs
        return WebPageProcessor()
    else:
        # For unidentified local files, try to read as text
        class GenericFileProcessor(DocumentProcessor):
            def process(self, source: str, content: Optional[bytes] = None) -> List[Dict[str, Any]]:
                try:
                    metadata = {"source": source, "source_type": "generic"}
                    if content:
                        # Try to decode as text
                        text = content.decode('utf-8', errors='replace')
                    else:
                        # Try to read as text file
                        with open(source, 'r', encoding='utf-8', errors='replace') as f:
                            text = f.read()
                    return self.chunk_text(text, metadata)
                except Exception as e:
                    print(f"Error processing generic file {source}: {e}")
                    return []
                    
        return GenericFileProcessor()
